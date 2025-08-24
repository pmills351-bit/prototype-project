# src/report_docx.py
from __future__ import annotations
from typing import Dict, Any, Optional, List
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import pandas as pd

MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# ---------- helpers ----------
def _h1(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run.font.size = Pt(18)
    run.bold = True
    return p

def _h2(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.bold = True
    return p

def _p(doc: Document, text: str):
    doc.add_paragraph(text)

def _table_from_dataframe(doc: Document, df: pd.DataFrame, max_rows: int = 400):
    df_show = df.copy()
    if len(df_show) > max_rows:
        df_show = df_show.head(max_rows)
    rows, cols = df_show.shape
    table = doc.add_table(rows=rows + 1, cols=cols)
    table.style = "Table Grid"
    # header
    for j, col in enumerate(df_show.columns):
        table.cell(0, j).text = str(col)
    # body
    for i in range(rows):
        for j in range(cols):
            val = df_show.iat[i, j]
            table.cell(i + 1, j).text = "" if pd.isna(val) else str(val)
    table.autofit = True
    return table

# ---------- single-run report ----------
def build_docx_report_bytes(
    *,
    app_build: str,
    settings: Dict[str, Any],
    table_df: pd.DataFrame,
    calibration_png_b64: Optional[str] = None,
    brier: Optional[float] = None,
) -> BytesIO:
    """Build the single-run DOCX and return a BytesIO buffer."""
    doc = Document()
    # margins
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    # Cover
    _h1(doc, "Bias & Equity Audit Report")
    _p(doc, app_build)
    doc.add_paragraph()

    # Executive Summary (simple counts)
    _h2(doc, "Executive Summary")
    parity_col = None
    for c in table_df.columns:
        if str(c).lower().startswith("parity"):
            parity_col = c
            break
    lines = []
    if parity_col and not table_df.empty:
        counts = table_df[parity_col].astype(str).value_counts()
        fails = int(counts.get("❌ Fail", 0) + counts.get("Fail", 0))
        border = int(counts.get("⚠️ Borderline", 0) + counts.get("Borderline", 0))
        ok = int(counts.get("✅ Pass", 0) + counts.get("Pass", 0))
        lines.append(f"Pass: {ok}, Borderline: {border}, Fail: {fails}")
    else:
        lines.append("No parity column found in table.")
    if brier is not None:
        lines.append(f"Brier score: {brier:.3f}")
    _p(doc, " • ".join(lines))
    doc.add_paragraph()

    # Settings
    _h2(doc, "Audit Settings")
    for k, v in settings.items():
        _p(doc, f"{k}: {v}")
    doc.add_paragraph()

    # Results table
    _h2(doc, "Group / Intersectional Summary")
    _table_from_dataframe(doc, table_df)

    # Calibration plot
    if calibration_png_b64:
        import base64, tempfile
        raw = base64.b64decode(calibration_png_b64)
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp.write(raw)
            tmp.flush()
            _h2(doc, "Calibration: Reliability Diagram")
            doc.add_picture(tmp.name, width=Inches(5.5))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def build_docx_report(
    *,
    path: str,
    app_build: str,
    settings: Dict[str, Any],
    table_df: pd.DataFrame,
    calibration_png_b64: Optional[str] = None,
    brier: Optional[float] = None,
) -> str:
    """Write single-run DOCX to disk and return path (compat)."""
    buf = build_docx_report_bytes(
        app_build=app_build,
        settings=settings,
        table_df=table_df,
        calibration_png_b64=calibration_png_b64,
        brier=brier,
    )
    with open(path, "wb") as f:
        f.write(buf.getvalue())
    return path

# ---------- two-run comparison report ----------
def build_docx_compare_bytes(
    *,
    app_build: str,
    runA_title: str,
    runB_title: str,
    settingsA: Dict[str, Any],
    settingsB: Dict[str, Any],
    key_cols: List[str],
    merged_df: pd.DataFrame,
) -> BytesIO:
    """
    Build a comparison DOCX from the 'merged' table produced in the Compare tab.

    merged_df should include columns:
      key_cols..., rate_A, rate_B, Δrate, disp_A, disp_B, Δdisp, parity_A, parity_B, parity_change
    """
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    # Cover
    _h1(doc, "Bias & Equity Audit — Comparison Report")
    _p(doc, app_build)
    doc.add_paragraph()

    # Overview
    _h2(doc, "Overview")
    _p(doc, f"Run A: {runA_title}")
    _p(doc, f"Run B: {runB_title}")
    doc.add_paragraph()

    # Executive Summary (parity change counts)
    _h2(doc, "Executive Summary")
    parity_change_col = "parity_change"
    lines = []
    if parity_change_col in merged_df.columns:
        counts = merged_df[parity_change_col].astype(str).value_counts()
        # try to identify interesting flips
        flips_up = 0
        flips_down = 0
        for v, cnt in counts.items():
            s = str(v)
            if "Pass →" in s or "✅ Pass →" in s:
                flips_down += 0  # no-op, track specific if needed
            if "→ ✅ Pass" in s:
                flips_up += cnt
            if "→ ⚠️ Borderline" in s or "→ Borderline" in s:
                pass  # could track separately
            if "→ ❌ Fail" in s or "→ Fail" in s:
                pass
        lines.append("Parity changes:")
        for v, cnt in counts.items():
            lines.append(f"  - {v}: {cnt}")
        if flips_up:
            lines.append(f"Improved to Pass: {flips_up}")
    else:
        lines.append("No parity_change column found.")
    _p(doc, "\n".join(lines))
    doc.add_paragraph()

    # Settings side-by-side (basic dump)
    _h2(doc, "Run Settings")
    _p(doc, "Run A settings:")
    for k, v in settingsA.items():
        _p(doc, f"  {k}: {v}")
    _p(doc, "Run B settings:")
    for k, v in settingsB.items():
        _p(doc, f"  {k}: {v}")
    doc.add_paragraph()

    # Results table
    _h2(doc, "Side-by-side Summary (A vs B)")
    show_cols = key_cols + ["rate_A","rate_B","Δrate","disp_A","disp_B","Δdisp","parity_A","parity_B","parity_change"]
    show_cols = [c for c in show_cols if c in merged_df.columns]
    _table_from_dataframe(doc, merged_df[show_cols])

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


